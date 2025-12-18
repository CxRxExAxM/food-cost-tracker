"""
File processing service for AI recipe parser.

Extracts text from uploaded documents (.docx, .pdf, .xlsx).
"""

import io
from typing import BinaryIO
from fastapi import UploadFile, HTTPException


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded document.

    Supports:
    - .docx (Microsoft Word)
    - .pdf (PDF documents)
    - .xlsx (Excel spreadsheets)

    Args:
        file: Uploaded file object

    Returns:
        Extracted text content

    Raises:
        HTTPException: If file format unsupported or extraction fails
    """

    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    filename_lower = file.filename.lower()

    try:
        # Read file content into memory
        content = await file.read()

        if filename_lower.endswith('.docx'):
            return await extract_from_docx(io.BytesIO(content))
        elif filename_lower.endswith('.pdf'):
            return await extract_from_pdf(io.BytesIO(content))
        elif filename_lower.endswith('.xlsx'):
            return await extract_from_excel(io.BytesIO(content))
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file format. Supported: .docx, .pdf, .xlsx"
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )


async def extract_from_docx(file_content: BinaryIO) -> str:
    """
    Extract text from Word document (.docx).

    Args:
        file_content: File content as binary stream

    Returns:
        Plain text content
    """

    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Install with: pip install python-docx"
        )

    try:
        doc = Document(file_content)

        # Extract all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(' | '.join(row_text))

        # Combine paragraphs and tables
        all_text = paragraphs + table_text

        if not all_text:
            raise HTTPException(
                status_code=400,
                detail="Document appears to be empty"
            )

        return '\n'.join(all_text)

    except Exception as e:
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(
            status_code=400,
            detail=f"Invalid or corrupted Word document: {str(e)}"
        )


async def extract_from_pdf(file_content: BinaryIO) -> str:
    """
    Extract text from PDF document.

    Note: This is basic text extraction. OCR for scanned PDFs
    is planned for a future phase.

    Args:
        file_content: File content as binary stream

    Returns:
        Plain text content
    """

    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="pypdf not installed. Install with: pip install pypdf"
        )

    try:
        reader = PdfReader(file_content)

        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(page_text)

        if not text_parts:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from PDF. This might be a scanned document (OCR support coming soon)"
            )

        return '\n\n'.join(text_parts)

    except Exception as e:
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(
            status_code=400,
            detail=f"Invalid or corrupted PDF document: {str(e)}"
        )


async def extract_from_excel(file_content: BinaryIO) -> str:
    """
    Extract text from Excel spreadsheet (.xlsx).

    Assumes recipe data is in first sheet with standard format:
    - First rows might be recipe name/info
    - Subsequent rows are ingredients

    Args:
        file_content: File content as binary stream

    Returns:
        Plain text content with formatting preserved
    """

    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="openpyxl not installed. Install with: pip install openpyxl"
        )

    try:
        workbook = load_workbook(file_content, data_only=True)
        sheet = workbook.active

        if not sheet:
            raise HTTPException(
                status_code=400,
                detail="Excel file has no active sheet"
            )

        # Extract all rows
        rows = []
        for row in sheet.iter_rows(values_only=True):
            # Filter out empty cells and convert to strings
            row_data = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if row_data:
                rows.append(' | '.join(row_data))

        if not rows:
            raise HTTPException(
                status_code=400,
                detail="Excel sheet appears to be empty"
            )

        return '\n'.join(rows)

    except Exception as e:
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(
            status_code=400,
            detail=f"Invalid or corrupted Excel document: {str(e)}"
        )


async def validate_file_before_parse(file: UploadFile) -> dict:
    """
    Pre-validate file before sending to AI.

    This quick check prevents wasting credits on obviously invalid files.

    Args:
        file: Uploaded file object

    Returns:
        Dict with validation results

    Raises:
        HTTPException: If file fails validation
    """

    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    # Check file extension
    filename_lower = file.filename.lower()
    if not any(filename_lower.endswith(ext) for ext in ['.docx', '.pdf', '.xlsx']):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Supported: .docx, .pdf, .xlsx"
        )

    # Check file size (10MB limit)
    content = await file.read()
    file_size = len(content)

    if file_size > 10 * 1024 * 1024:  # 10MB
        raise HTTPException(
            status_code=400,
            detail="File size exceeds 10MB limit"
        )

    if file_size == 0:
        raise HTTPException(
            status_code=400,
            detail="File is empty"
        )

    # Reset file pointer for later reading
    await file.seek(0)

    # Try to extract text (quick validation)
    try:
        text = await extract_text_from_file(file)

        # Check minimum text length (recipes should have at least some content)
        if len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Document appears to have insufficient content for a recipe"
            )

        # Reset file pointer again
        await file.seek(0)

        return {
            'valid': True,
            'filename': file.filename,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'text_length': len(text),
            'message': 'File is valid for parsing'
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Could not read file: {str(e)}"
        )
